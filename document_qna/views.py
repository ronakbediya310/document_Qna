import os
import json
import time
import pymysql
import re
import pickle
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse, FileResponse
from django.shortcuts import get_object_or_404
from .models import SourceDocument
from .models import UnansweredQuestion
from django.core.files import File
from django.core.cache import cache
from docx import Document
from django.http import JsonResponse
from django.shortcuts import render
from django.core.files.storage import default_storage
from langchain_groq import ChatGroq
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyPDFLoader, CSVLoader, JSONLoader
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import LLMChain
from langchain.chains.combine_documents import create_stuff_documents_chain
from dotenv import load_dotenv
from langchain.memory import ConversationBufferWindowMemory
from langchain.chains import ConversationalRetrievalChain
from django.utils.timezone import now
from datetime import timedelta
from rest_framework.response import Response
from rest_framework.decorators import api_view
import uuid
from django.conf import settings
from .models import APIToken,AdminUser
from django.utils.timezone import now


## Initialization
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
llm = ChatGroq(groq_api_key=GROQ_API_KEY, model="llama-3.3-70b-versatile")

VECTOR_STORE = None
FAISS_INDEX_PATH = "faiss_index.pkl"

MYSQL_CONFIG = {
    "host": os.getenv("DB_HOST"),
    "user": os.getenv("DB_USER"),
    "password": os.getenv("DB_PASSWORD"),
    "database": os.getenv("DB_NAME"),
    "port": int(os.getenv("DB_PORT", 3306)),
}
## Hugging face embedding model to convert text chunks into numerical vectors.
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

prompt = ChatPromptTemplate.from_template(
    """
    You are an AI assistant specialized in answering customer queries.  
    Provide accurate, in-depth, and concise responses strictly based on the given context and Chat History.  

    **Guidelines:**  
    - Your response must be **precise and to the point**.  
    - **Do not** use phrases like "According to the provided context."  
    - **Never** disclose API keys, passwords, or personal user information.  
    - **If the answer is not in the context, respond with:** "Not enough info to answer." 
    Chat History:
    {chat_history}

    <context>
    {context}
    </context>
     question: {question} 
    """
)

latest_query = ""
## Memory concept :
CHAT_HISTORY_LIMIT = 10
memory = ConversationBufferWindowMemory(
    memory_key="chat_history",
    output_key="answer",
    k=CHAT_HISTORY_LIMIT,
    return_messages=True,
)


def get_chat_history(memory):
    memory.load_memory_variables({}).get("chat_history", [])


def clear_chat_history():
    """Clear chat memory."""
    memory.clear()
    print("Chat history cleared.")


## Database operations:
def connect_db():
    return pymysql.connect(**MYSQL_CONFIG, cursorclass=pymysql.cursors.DictCursor)


## Storing embedings into DB.
def store_vector_in_db(file_name, chunk_text, embedding):
    conn = connect_db()
    try:
        with conn.cursor() as cursor:
            sql = "INSERT INTO vector_store (file_name, chunk_text, embedding) VALUES (%s, %s, %s)"
            cursor.execute(sql, (file_name, chunk_text, json.dumps(embedding)))
        conn.commit()
    except Exception as e:
        print(f"Error storing vector: {e}")
        conn.rollback()
    finally:
        conn.close()


## Loading chunks,and vector Embeddings from DB/
def load_vectors_from_db():
    """Load vectors from MySQL."""
    global VECTOR_STORE
    conn = connect_db()
    try:
        with conn.cursor() as cursor:
            cursor.execute("SELECT chunk_text, embedding FROM vector_store LIMIT 10000")
            data = cursor.fetchall()
        ## DB is Empty
        if not data:
            print("No vector data found in the database.")
            return

        documents = []
        embeddings_list = []

        for d in data:
            try:
                chunk_text = d["chunk_text"]
                embedding = json.loads(d["embedding"])

                if not isinstance(embedding, list):
                    print(f"Invalid embedding format: {embedding}")
                    continue

                documents.append(chunk_text)
                embeddings_list.append(embedding)
            except Exception as e:
                print(f"Error processing embedding: {e}")

        print(f"Loaded {len(embeddings_list)} embeddings successfully.")
        if len(embeddings_list) != len(documents):
            print("Warning: Mismatch between documents and embeddings!")
        VECTOR_STORE = FAISS.from_texts(documents, embeddings)
        save_faiss_index()

    except Exception as e:
        print(f"Error loading vectors: {e}")
    finally:
        conn.close()


def save_faiss_index():
    """Save the FAISS index to a Pickle file."""
    global VECTOR_STORE
    if VECTOR_STORE is not None:
        with open(FAISS_INDEX_PATH, "wb") as f:
            pickle.dump(VECTOR_STORE, f)
        print("FAISS index saved successfully.")


## Index file for optimized data retrievals.
def load_faiss_index():
    """Load the FAISS index from Pickle if available."""
    global VECTOR_STORE
    if os.path.exists(FAISS_INDEX_PATH):
        with open(FAISS_INDEX_PATH, "rb") as f:
            VECTOR_STORE = pickle.load(f)
        print("FAISS index loaded from Pickle.")
        return True
    return False


## This function is used to Retrieve vector embeddings either from index file or MySQL DB.
def load_index_or_db():
    """Ensures FAISS index is loaded, otherwise loads from DB."""
    global VECTOR_STORE
    if not load_faiss_index():
        load_vectors_from_db()
    return VECTOR_STORE is not None


## Extracting information from JSON Data.
def extract_json_text_as_chunks(data):
    """Convert each JSON object into a single chunk."""
    if isinstance(data, list):
        return [json.dumps(item, ensure_ascii=False) for item in data]
    elif isinstance(data, dict):
        return [json.dumps(data, ensure_ascii=False)]
    print("Unexpected JSON format: Root element is not a dict or list.")
    return []


## This function triggers when user uploads new document.
def process_new_document(file_path):
    try:
        if not os.path.exists(file_path):
            print(f"File not found: {file_path}")
            return

        text_entries = []

        if file_path.endswith(".pdf"):
            loader = PyPDFLoader(file_path)
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000, chunk_overlap=200
            )
            documents = loader.load()
            text_entries = text_splitter.split_documents(documents)
        elif file_path.endswith(".csv"):
            text_entries = [
                doc.page_content for doc in CSVLoader(file_path=file_path).load()
            ]
        elif file_path.endswith(".json"):
            with open(file_path, "r", encoding="utf-8") as f:
                json_data = json.load(f)
            if json_data:
                text_entries = extract_json_text_as_chunks(json_data)
            else:
                print(f"JSON file is empty: {file_path}")
                return
        else:
            print(f"Unsupported file format: {file_path}")
            return

        if not text_entries:
            print(f"No valid text extracted from: {file_path}")
            return

        for text in text_entries:
            embedding = embeddings.embed_query(text)
            store_vector_in_db(os.path.basename(file_path), text, embedding)

        load_vectors_from_db()
        ## Remove file from Media folder.
        os.remove(file_path)
        print(f"File processed and deleted: {file_path}")
        return True
    except Exception as e:
        print(f"Error processing document ({file_path}): {e}")
        return False


def save_unanswered_question(question):
    """Store unanswered questions in the database."""
    if not UnansweredQuestion.objects.filter(question=question).exists():
        UnansweredQuestion.objects.create(question=question)


def get_unanswered_questions():
    return UnansweredQuestion.objects.all()


# Global variable to store retrieval_chain in memory
retrieval_chain_instance = None


def get_retrieval_chain():
    global retrieval_chain_instance

    if retrieval_chain_instance is not None:
        return retrieval_chain_instance

    if not load_index_or_db():
        return None

    retriever = cache.get("retriever")
    if retriever is None:
        retriever = VECTOR_STORE.as_retriever(search_kwargs={"k": 5})
        cache.set("retriever", retriever, timeout=3600)

    retrieval_chain_instance = ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=retriever,
        combine_docs_chain_kwargs={"prompt": prompt},
        memory=memory,
        return_source_documents=True,
    )
    return retrieval_chain_instance


def store_retrieved_documents(user_query, response, save_results=False):
    """
    Stores each retrieved document separately in the database while keeping its associated source type.
    """
    if not save_results:
        return None

    folder_name = "media/retrieved_docs"
    os.makedirs(folder_name, exist_ok=True)

    safe_filename = re.sub(r'[\\/*?:"<>|]', "_", user_query)

    for i, content in enumerate(response.get("source_documents", []), start=1):
        doc_filename = os.path.join(folder_name, f"{safe_filename}_doc{i}.docx")

        # Extract source type & title
        source_type = "unknown"
        extracted_title = f"Document_{i}"  # Default title if not found
        dynamic_url = ""

        try:
            doc_data = json.loads(content.page_content)

            if "id" in doc_data:
                source_type = "article"
                dynamic_url = (
                    f"https://manage.accuwebhosting.com/knowledgebase/{doc_data['id']}"
                )
            elif "tid" in doc_data:
                source_type = "ticket"
                dynamic_url = f"https://manage.accuwebhosting.com/admin20112012/supporttickets.php?action=viewticket&id={doc_data['tid']}"
            elif "internal_id" in doc_data:
                source_type = "internal_wiki"
                dynamic_url = f"https://support.jilesh.com/internal_wiki/{doc_data['internal_id']}"
            elif "website_url" in doc_data:
                source_type = "website"
                dynamic_url = doc_data["website_url"]

            # Extract title/question/query
            extracted_title = (
                doc_data.get("title")
                or doc_data.get("question")
                or doc_data.get("query")
                or extracted_title
            )

        except Exception as e:
            print(f"Error parsing document content: {e}")

        # Save document with extracted title
        # doc = Document()
        # doc.add_heading(f"Results for: {extracted_title}", level=1)
        # doc.add_paragraph(content.page_content)
        # doc.save(doc_filename)

        # Store each document separately
        source_doc = SourceDocument.objects.create(
            query=user_query,
            title=extracted_title,
            file_path=dynamic_url,
            content=content.page_content,
            source=source_type,
        )

        print(
            f" (Title: {extracted_title}, Source Type: {source_type}, URL: {dynamic_url})"
        )

    return True


def view_source_document(query):
    """
    Fetches all source documents for a given query and returns them as a list.
    """
    documents = SourceDocument.objects.filter(query=query)

    if not documents.exists():
        return JsonResponse({"error": "No documents found for this query"}, status=404)

    # Collect all document contents along with their source types and URLs
    results = [
        {
            "title": doc.title,
            "source": doc.source,
            "file_path": doc.file_path,
        }
        for doc in documents
    ]
    return JsonResponse({"documents": results})


def download_source_document(query):
    document = get_object_or_404(SourceDocument, query=query)

    if document.file_path:
        return FileResponse(open(document.file_path.path, "rb"), as_attachment=True)
    else:
        return JsonResponse({"error": "No file available"}, status=404)


def get_response(user_query):
    """Handles the query processing logic."""
    retrieval_chain = get_retrieval_chain()
    if retrieval_chain is None:
        return {"error": "No indexed documents found. Please upload some files first."}, 400

    start_time = time.process_time()
    response = retrieval_chain.invoke({"question": user_query})
    response_time = time.process_time() - start_time

    answer = response.get("answer", "").strip()
    if not answer or any(
        phrase in answer.lower()
        for phrase in [
            "i don't know",
            "i'm not sure",
            "i don't have enough information",
            "no relevant information",
            "Not enough info to answer.",
        ]
    ):
        save_unanswered_question(user_query)

    store_retrieved_documents(user_query, response, True)
    formatted_answer = format_response(answer)

    return JsonResponse(
        {"answer": formatted_answer, "response_time": response_time},safe=False
    ) 


def format_response(response):
    """Formats the bot's response beautifully with better readability."""
    
    # Convert markdown-like formatting to HTML-friendly format
    response = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", response) 
    response = re.sub(r"\*(.*?)\*", r"<i>\1</i>", response)  
    
    # Ensure new lines before bullet points and numbered lists for spacing
    response = re.sub(r"(\n?[-•*]\s+)", r"<br>\1", response)  # Unordered lists
    response = re.sub(r"(\n?\d+\.\s+)", r"<br>\1", response)  
    
    # Replace double newlines with paragraph breaks for better spacing
    response = re.sub(r"\n{2,}", r"<br><br>", response)

    return response.strip()

def create_dummy_api_token():
    """Function to create a dummy user and API token."""
    user, created = AdminUser.objects.get_or_create(
        username="Ronak",
        
    )

    api_token = APIToken.objects.create(
        user=user,
        token=str(uuid.uuid4()),  
        expires_at=now() + timedelta(days=7)  
    )

    print(f"Dummy token created: {api_token.token}")
    return api_token.token


def get_user_from_token(token):
    """Validate user and token by comparing request token and DB token."""
    try:
        user = AdminUser.objects.get(username="Ronak") 
        api_token = APIToken.objects.select_related("user").get(user=user, token=token)
        if api_token.is_valid():
            return api_token.user
    except APIToken.DoesNotExist:
        return None
    return None



# ## This is the Index view of Django.
def index(request):
    # create_dummy_api_token()
    if request.method == "POST":
        ## Document Upload section.
        if "document" in request.FILES:
            file = request.FILES["document"]
            file_path = default_storage.save(f"media/{file.name}", file)
            if process_new_document(file_path):
                return render(request, "index.html", {"message": "Upload successful!"})
            return render(request, "index.html", {"error": "Upload failed!"})

        elif "fetch" in request.POST and "query" in request.POST:
            return view_source_document(query=request.POST["query"])

        elif "query" in request.POST:
            return get_response(request.POST["query"])

    return render(request, "index.html")
@csrf_exempt
@api_view(["POST","GET"])
def query_api(request):
    """REST API to validate token and process queries."""
    
    # Ensure JSON data is received
    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return Response({"error": "Invalid JSON format"}, status=400)

    token = data.get("token") or  request.headers.get("Authorization")
    user_query = data.get("query")

    if not token or not user_query:
        return Response({"error": "Token and query are required."}, status=400)

    user = get_user_from_token(token)
    if not user:
        return Response({"error": "Invalid or expired token."}, status=403)

    return get_response(user_query) 